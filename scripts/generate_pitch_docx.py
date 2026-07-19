from pathlib import Path
from typing import Optional
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "public" / "SurgiFind-Pitch-Deck.docx"
ARCHITECTURE_IMAGE = ROOT / "public" / "long-run-architecture.svg"


def add_title(document: Document, text: str, subtitle: Optional[str] = None) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(24)
    if subtitle:
        s = document.add_paragraph()
        s.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sr = s.add_run(subtitle)
        sr.italic = True
        sr.font.size = Pt(11)


def add_heading(document: Document, text: str) -> None:
    document.add_heading(text, level=1)


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        document.add_paragraph(item, style="List Bullet")


def add_numbered(document: Document, items: list[str]) -> None:
    for item in items:
        document.add_paragraph(item, style="List Number")


def add_note(document: Document, text: str) -> None:
    p = document.add_paragraph()
    label = p.add_run("Presenter note: ")
    label.bold = True
    p.add_run(text)


def main() -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    add_title(
        document,
        "SurgiFind Pitch Deck",
        "Shareable hackathon pitch + future-state architecture overview",
    )

    add_heading(document, "1. SurgiFind")
    document.add_paragraph(
        "Find the right hospital for surgery with clarity on price, insurance, and booking."
    )
    document.add_paragraph(
        "SurgiFind helps patients discover surgical care faster by combining hospital search, price comparison, insurance matching, and guided booking in one platform."
    )

    add_heading(document, "2. Problem")
    add_bullets(document, [
        "Hospital discovery is fragmented across search engines, aggregators, and hospital websites.",
        "Price transparency is poor, so families cannot compare likely costs with confidence.",
        "Insurance network and coverage checks are confusing and time-consuming.",
        "Booking a consultation or surgery often requires manual calls and follow-ups.",
    ])
    add_note(document, "Patients need clarity, comparison, and a trusted next step instead of fragmented search.")

    add_heading(document, "3. Solution")
    add_bullets(document, [
        "Hospital and surgery search",
        "Cost range comparison",
        "Insurance plan matching",
        "Conversational guidance",
        "Authenticated booking and booking history",
    ])

    add_heading(document, "4. Product Demo Flow")
    add_numbered(document, [
        "Search by surgery, city, budget, rating, or hospital type.",
        "Compare hospitals with price ranges and availability.",
        "Use the chat assistant in natural language.",
        "Review matching insurance plans.",
        "Book a consultation or surgery slot securely.",
        "Track or cancel the booking from booking history.",
    ])

    add_heading(document, "5. Key Features")
    add_bullets(document, [
        "Smart surgery search with filters",
        "Conversational assistant with direct search and symptom-assisted guidance",
        "Insurance matching against covered surgeries and network hospitals",
        "Secure authentication with booking history",
        "Slot reservation flow with cancellation support",
    ])

    add_heading(document, "6. Differentiation")
    add_bullets(document, [
        "Connected workflow from search to booking, not just content browsing",
        "Combines search, guidance, insurance relevance, and booking execution",
        "Patient-first decision support instead of isolated hospital listings",
    ])

    add_heading(document, "7. Technology and Trust")
    add_bullets(document, [
        "Frontend: Next.js, React, TypeScript, Tailwind CSS",
        "Backend: Next.js route handlers",
        "Data and identity: Supabase Auth and PostgreSQL",
        "AI assistant: OpenAI-backed conversational guidance with fallbacks",
        "Security baseline: server-side writes, authenticated booking flow, row-level access controls",
    ])

    add_heading(document, "8. Business Potential")
    add_bullets(document, [
        "Lead generation fees from hospitals",
        "Premium listings or sponsored discovery placement",
        "Booking conversion commissions",
        "Insurance partnership and referral workflows",
        "Care navigation subscriptions for high-intent users",
    ])

    add_heading(document, "9. Go-To-Market")
    add_bullets(document, [
        "Start with high-consideration elective surgeries",
        "Target urban patients comparing providers across price bands",
        "Partner with hospitals seeking qualified inbound demand",
        "Use SEO, conversational search, and referrals to capture intent",
    ])

    add_heading(document, "10. Next Steps")
    add_bullets(document, [
        "Add a payment system for deposits, booking confirmation, refunds, and receipts",
        "Complete production cloud deployment with CI/CD, secrets handling, and monitoring",
        "Move booking reservation and insert into stronger transactional database workflows",
        "Add doctor-level scheduling and richer hospital-side operations",
        "Introduce rate limiting, WAF, audit trails, and production security hardening",
        "Expand hospital and insurance integrations",
    ])

    add_heading(document, "11. Long-Run Architecture")
    document.add_paragraph(
        "Future-state production platform built around the current app with edge delivery, load balancing, API management, cache, autoscaling runtime, observability, and payment infrastructure."
    )
    add_bullets(document, [
        "Global edge delivery and CDN",
        "Load balancer and autoscaled application runtime",
        "API management for traffic policies and versioned APIs",
        "Cache layer for catalog search and frequent reads",
        "Background jobs for notifications, reconciliation, and reporting",
        "Security controls including WAF, secrets management, and audit logging",
    ])
    document.add_paragraph(
        "Architecture diagram file included separately in the repository at public/long-run-architecture.svg."
    )

    add_heading(document, "12. Team")
    add_bullets(document, [
        "Aditi Prasad",
        "Mayank Prasad",
        "Deepnarayan Lohra",
    ])

    add_heading(document, "13. Closing")
    document.add_paragraph(
        "When patients need surgery, they should not have to navigate cost, trust, insurance, and booking alone. SurgiFind gives them one place to decide and act."
    )
    add_bullets(document, [
        "Pilot hospital partners",
        "Product feedback",
        "Healthcare domain mentors",
        "Early users for usability testing",
    ])

    add_heading(document, "Appendix: Suggested Production Deployment")
    add_bullets(document, [
        "Frontend and server routes on Vercel",
        "Supabase for auth and primary transactional data",
        "Managed Redis for caching and rate limits",
        "API gateway or edge controls for request governance",
        "Monitoring stack for logs, uptime, traces, and alerting",
    ])

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    main()
